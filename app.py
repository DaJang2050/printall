# PrintALLApp.py
import os
import sys
import glob
import shutil
import logging
import threading
import tempfile
import subprocess
import tkinter as tk
from logging.handlers import RotatingFileHandler
from tkinter import ttk, filedialog, messagebox, scrolledtext
import webbrowser

# --- 统一的依赖项导入 ---

# 通用
from PIL import Image, ImageDraw, ImageFont, ImageStat

# 水印功能
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor,Pt

# <--- 使用 PyMuPDF (fitz) 作为新的PDF水印引擎 ---
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
# <--- END ---

# <--- 添加对Excel文件的支持 ---
try:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Side
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
# <--- END ---

# 打印功能
from pypdf import (
    PdfWriter as PrintPdfWriter,
    PdfReader as PrintPdfReader,
)  # 使用别名避免冲突

# 打印功能 - Windows特定
try:
    import win32api
    import functools
    import win32print
    import ctypes
    _StrCmpLogicalW = ctypes.windll.shlwapi.StrCmpLogicalW
    PYWIN32_AVAILABLE = True
except (ImportError, AttributeError):
    PYWIN32_AVAILABLE = False
    _StrCmpLogicalW = None
    win32print = None
    functools = None


# --- 全局配置 ---

# [打印功能配置] 请根据您的系统修改此路径, 它将作为默认值
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

# [水印功能配置] 使用 Windows 系统字体路径
CHINESE_FONT_PATH = "C:/Windows/Fonts/msyh.ttc"

# [打印功能配置]
PRINT_DPI = 300
A4_WIDTH_MM = 210
A4_HEIGHT_MM = 297

def resource_path(relative_path):
    """ 获取资源的绝对路径, 适用于开发环境和 Nuitka/PyInstaller 打包环境 """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- 检查水印字体和依赖 ---
# 直接检查绝对路径，因为仅在windows下运行
CHINESE_FONT_AVAILABLE = os.path.exists(CHINESE_FONT_PATH)
if not CHINESE_FONT_AVAILABLE:
    print(f"警告：PDF水印所需的中文字体文件 '{CHINESE_FONT_PATH}' 未在系统中找到。")
    print("PDF水印中的中文可能无法显示。")


class PrintALLApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PrintALL 全能打印助手  |  By YangZhiqiang")

        try:
            icon_path = resource_path("PrintALL.ico") 
            self.root.iconbitmap(icon_path)
        except Exception as e:
            print(f"无法设置窗口图标: {e}")
        
        self._setup_logging()
        self.logger.info("========================================")
        self.logger.info("应用程序启动")
        self.logger.info(f"pywin32可用 (打印支持): {PYWIN32_AVAILABLE}")
        self.logger.info(f"PyMuPDF可用 (PDF水印支持): {PYMUPDF_AVAILABLE}")
        self.logger.info(f"中文字体文件可用: {CHINESE_FONT_AVAILABLE}")
        self.logger.info(f"openpyxl可用 (Excel支持): {OPENPYXL_AVAILABLE}")

        window_width = 800
        window_height = 600
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        center_x = int((screen_width - window_width) / 2)
        center_y = int((screen_height - window_height) / 2)
        self.root.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(pady=10, padx=10, fill="both", expand=True)

        self.watermark_tab = ttk.Frame(self.notebook, padding="10")
        self.print_tab = ttk.Frame(self.notebook, padding="10")

        self.notebook.add(self.watermark_tab, text="添加水印")
        self.notebook.add(self.print_tab, text="批量打印")

        self._initialize_watermark_vars()
        self._setup_watermark_tab()

        self._initialize_print_vars()
        self._setup_print_tab()

        self.log_watermark("注意：本工具会直接修改原始文件，请务必提前备份重要数据！")
        if not PYMUPDF_AVAILABLE:
            self.log_watermark(
                "【警告】'PyMuPDF' 模块缺失，无法处理 .pdf 文件。请运行 'pip install PyMuPDF' 安装。"
            )
        if not CHINESE_FONT_AVAILABLE:
            self.log_watermark(
                f"【警告】中文字体 '{CHINESE_FONT_PATH}' 未找到，PDF水印中的中文可能无法正常显示。"
            )
        if not OPENPYXL_AVAILABLE:
            self.log_watermark(
                "【警告】'openpyxl' 模块缺失，无法处理 .xlsx 文件。请运行 'pip install openpyxl' 安装。"
            )

        self._initialize_print_log()
        self.root.after(100, self._start_background_tasks)


    def _setup_logging(self):
        log_dir = os.path.dirname(os.path.abspath(__file__))
        log_file = os.path.join(log_dir, "PrintALL.log")
        self.logger = logging.getLogger("PrintALLAppLogger")
        self.logger.setLevel(logging.INFO)
        if self.logger.hasHandlers():
            self.logger.handlers.clear()
        handler = RotatingFileHandler(
            log_file, maxBytes=5 * 1024 * 1024, backupCount=3, encoding="utf-8"
        )
        formatter = logging.Formatter(
            "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
        )
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)

    def on_closing(self):
        self.logger.info("应用程序关闭")
        self.logger.info("========================================\n")
        self.root.destroy()
        
    # ==================================================================
    # I. 水印功能模块
    # ==================================================================
    def _initialize_watermark_vars(self):
        self.watermark_folder_path = tk.StringVar()
        self.process_word = tk.BooleanVar(value=True)
        self.process_pic = tk.BooleanVar(value=True)
        self.process_pdf = tk.BooleanVar(value=True)
        self.process_excel = tk.BooleanVar(value=True)
        self.pic_opacity = tk.IntVar(value=150)
        self.pic_position = tk.StringVar(value="顶部居中")

    def _validate_entry(self, new_value, min_val, max_val):
        if new_value == "":
            return True
        try:
            value = int(new_value)
            return min_val <= value <= max_val
        except ValueError:
            return False

    def _setup_watermark_tab(self):
        folder_frame = ttk.LabelFrame(
            self.watermark_tab, text="第一步: 选择要处理的文件夹", padding="10"
        )
        folder_frame.pack(fill=tk.X, pady=5)
        ttk.Entry(folder_frame, textvariable=self.watermark_folder_path, width=70).pack(
            side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5)
        )
        ttk.Button(
            folder_frame, text="浏览...", command=self._select_watermark_folder
        ).pack(side=tk.LEFT)

        type_frame = ttk.LabelFrame(self.watermark_tab, text="第二步: 选择要添加水印的文件类型", padding="10")
        type_frame.pack(fill=tk.X, pady=10)
        
        ttk.Checkbutton(type_frame, text="Word 文档 (.docx)", variable=self.process_word).pack(side=tk.LEFT, padx=10, pady=5)
        
        excel_checkbutton = ttk.Checkbutton(type_frame, text="Excel 工作簿 (.xlsx)", variable=self.process_excel)
        excel_checkbutton.pack(side=tk.LEFT, padx=10, pady=5)
        if not OPENPYXL_AVAILABLE:
            excel_checkbutton.config(state=tk.DISABLED)
        
        pdf_checkbutton = ttk.Checkbutton(type_frame, text="PDF 文档 (.pdf)", variable=self.process_pdf)
        pdf_checkbutton.pack(side=tk.LEFT, padx=10, pady=5)
        if not PYMUPDF_AVAILABLE or not CHINESE_FONT_AVAILABLE:
            pdf_checkbutton.config(state=tk.DISABLED)

        ttk.Checkbutton(type_frame, text="图片文件 (.jpg, .png, .bmp)", variable=self.process_pic).pack(side=tk.LEFT, padx=10, pady=5)

        settings_frame = ttk.LabelFrame(
            self.watermark_tab, text="第三步: 水印详细设置 (主要用于图片)", padding="10"
        )
        settings_frame.pack(fill=tk.X, pady=5)
        
        vcmd_opacity = (self.root.register(lambda p: self._validate_entry(p, 0, 255)), "%P",)
        
        # 调整了grid的column索引
        ttk.Label(settings_frame, text="透明度 (0-255):").grid(row=0, column=0, sticky=tk.W, padx=(5, 2), pady=5)
        ttk.Entry(settings_frame, textvariable=self.pic_opacity, width=8, validate="key", validatecommand=vcmd_opacity).grid(row=0, column=1, sticky=tk.W)
        ttk.Label(settings_frame, text="水印位置:").grid(row=0, column=2, sticky=tk.W, padx=(20, 2), pady=5)
        position_combo = ttk.Combobox(settings_frame, textvariable=self.pic_position, values=["左上角","右上角","左下角","右下角","居中","顶部居中","底部居中",],width=12,)
        position_combo.grid(row=0, column=3, sticky=tk.W, padx=5)
        position_combo.state(["readonly"])

        style = ttk.Style()
        style.configure("Accent.TButton", font=("Helvetica", 12, "bold"))
        self.watermark_process_button = ttk.Button(self.watermark_tab, text="批量添加水印", command=self.start_watermark_processing, style="Accent.TButton")
        self.watermark_process_button.pack(pady=20, ipady=5, fill="x", padx=5)
        
        log_frame = ttk.LabelFrame(self.watermark_tab, text="处理日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.watermark_log_text = scrolledtext.ScrolledText(log_frame, width=80, height=15, wrap=tk.WORD, font=("Consolas", 10))
        self.watermark_log_text.pack(fill=tk.BOTH, expand=True)

    def log_watermark(self, message):
        self.logger.info(f"[Watermark Tab] {message}")
        self.watermark_log_text.insert(tk.END, message + "\n")
        self.watermark_log_text.see(tk.END)
        self.root.update_idletasks()

    def _select_watermark_folder(self):
        folder = filedialog.askdirectory(title="选择要处理的文件夹")
        if folder:
            self.watermark_folder_path.set(folder)
            self.log_watermark(f"已选择文件夹: {folder}")
        
    def add_word_watermark(self, filepath):
        try:
            document = Document(filepath)
            filename = os.path.basename(filepath)
            header_text = f"打印对象：{filename}"
            
            # 获取文档的第一个节（section）来访问页面设置
            section = document.sections[0]
            header = section.header
            
            # 用于存储我们要操作的页眉段落
            target_paragraph = None
            
            # (检查页眉是否为空的逻辑保持不变)
            if not header.paragraphs:
                target_paragraph = header.add_paragraph()
            else:
                first_paragraph = header.paragraphs[0]
                if not first_paragraph.text.strip():
                    target_paragraph = first_paragraph
            
            if target_paragraph:
                target_paragraph.clear() 
                
                run = target_paragraph.add_run(header_text)
                
                # --- 核心修改部分：动态计算字体大小 ---
                
                # 1. 获取页面的可用宽度（页面宽度 - 左右页边距），单位转换为磅(Pt)
                #    section.page_width, left_margin, right_margin 的单位是 EMU，使用 .pt 可以直接转换为磅
                usable_width_pt = section.page_width.pt - section.left_margin.pt - section.right_margin.pt
                
                # 2. 根据可用宽度动态计算字体大小
                #    除数 60 是一个经验值，可以根据实际效果调整，值越大字体越小
                #    max(9, ...) 确保字体最小不低于9磅，保证可读性
                dynamic_fontsize = max(9, int(usable_width_pt / 60))
                
                # --- 修改结束 ---
                # 设置字体样式
                font = run.font
                font.name = '宋体'
                # 3. 应用动态计算出的字体大小
                font.size = Pt(dynamic_fontsize)
                font.color.rgb = RGBColor(255, 0, 0)
                
                # 设置段落对齐方式
                target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                document.save(filepath)
                self.log_watermark(f"[Word] ✔ 成功: {filename}")
            else:
                # 当 target_paragraph 为 None 时，说明页眉已存在且有内容，我们跳过了修改
                # 在这里打印“跳过”的日志
                self.log_watermark(f"[Word] ！ 跳过（页眉内容已存在）: {filename}")
            return True
        except Exception as e:
            self.log_watermark(f"[Word] ❌ 失败: {filename} - {e}")
            self.logger().error(f"处理Word文件'{filepath}'失败.", exc_info=True)
            return False
    
    def add_excel_watermark(self, filepath):
        try:
            filename = os.path.basename(filepath)
            workbook = openpyxl.load_workbook(filepath)
            header_format_string = "打印对象：&F|&A 第&[Page]/&N页"

            for ws in workbook.worksheets:
                # 判断Sheet是否为空
                is_sheet_none = (ws.max_row == 1 and ws.max_column == 1 and ws['A1'].value is None)
                # 判断工作表标签是否未设置颜色
                is_color_unset = (ws.sheet_properties.tabColor is None)

                if not is_sheet_none and len(workbook.worksheets) > 1 and is_color_unset:
                    ws.sheet_properties.tabColor = "d86100"
                if not ws.oddHeader.center.text:
                    ws.oddHeader.center.text = header_format_string
                ws.print_title_rows = '1:1'
                thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
                max_row = ws.max_row
                max_col = ws.max_column
                for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                    for cell in row:
                        if cell.value is not None:
                            cell.border = thin_border
                        if cell.row == 1:
                            cell.alignment = Alignment(horizontal='center', vertical='center')
            
            workbook.save(filepath)
            self.log_watermark(f"[Excel] ✔ 成功: {filename}")
            return True
        except Exception as e:
            self.log_watermark(f"[Excel] ❌ 失败: {filename} - {e}")
            self.logger.error(f"处理Excel文件'{filepath}'失败.", exc_info=True)
            return False
    
    def _get_pic_watermark_position(self, img_size, text_size):
        img_width, img_height = img_size
        text_width, text_height = text_size
        margin, top_bottom_margin = 20, 0
        positions = {
            "左上角": (margin, margin),
            "右上角": (img_width - text_width - margin, margin),
            "左下角": (margin, img_height - text_height - margin),
            "右下角": (img_width - text_width - margin, img_height - text_height - margin),
            "居中": ((img_width - text_width) // 2, (img_height - text_height) // 2),
            "顶部居中": ((img_width - text_width) // 2, top_bottom_margin),
            "底部居中": ((img_width - text_width) // 2, img_height - text_height - top_bottom_margin),
        }
        return positions.get(self.pic_position.get(), positions["顶部居中"])

    # 动态计算字体大小
    def add_picture_watermark(self, image_path):
        try:
            with Image.open(image_path) as img:
                original_format = img.format
                img_for_analysis = (img.convert("RGB") if img.mode not in ["RGB", "RGBA"] else img.copy())
                
                # 动态计算字体大小
                image_width = img.size[0]
                dynamic_font_size = max(25, int(image_width / 50))
                
                try:
                    # 直接使用 CHINESE_FONT_PATH
                    font = ImageFont.truetype(CHINESE_FONT_PATH, dynamic_font_size)
                except IOError:
                    font = ImageFont.load_default()
                    self.logger.warning(f"无法加载字体'{CHINESE_FONT_PATH}', 回退到默认字体。")
                
                watermark_text = f"打印对象：{os.path.basename(image_path)}"
                temp_draw = ImageDraw.Draw(Image.new("RGB", (1, 1)))
                bbox = temp_draw.textbbox((0, 0), watermark_text, font=font)
                text_width, text_height = bbox[2] - bbox[0], bbox[3] - bbox[1]
                x, y = self._get_pic_watermark_position(img.size, (text_width, text_height))
                stat = ImageStat.Stat(img_for_analysis.crop((x, y, x + text_width, y + text_height)))
                brightness = (0.299 * stat.mean[0] + 0.587 * stat.mean[1] + 0.114 * stat.mean[2])
                opacity_val = self.pic_opacity.get()
                fill_color = ((0, 0, 0, opacity_val) if brightness > 100 else (255, 255, 255, opacity_val))
                outline_color = ((255, 255, 255, opacity_val) if fill_color[0] == 0 else (0, 0, 0, opacity_val))
                if img.mode != "RGBA":
                    img = img.convert("RGBA")
                watermark_layer = Image.new("RGBA", img.size, (255, 255, 255, 0))
                draw = ImageDraw.Draw(watermark_layer)
                for dx in [-1, 0, 1]:
                    for dy in [-1, 0, 1]:
                        if dx != 0 or dy != 0:
                            draw.text((x + dx, y + dy), watermark_text, font=font, fill=outline_color)
                draw.text((x, y), watermark_text, font=font, fill=fill_color)
                watermarked = Image.alpha_composite(img, watermark_layer)
                if original_format in ["PNG", "BMP"]:
                    watermarked.save(image_path, format=original_format)
                else:
                    watermarked.convert("RGB").save(image_path, format="JPEG", quality=95)
                self.log_watermark(f"[图片] ✔ 成功: {os.path.basename(image_path)}")
                return True
        except Exception as e:
            self.log_watermark(f"[图片] ❌ 失败: {os.path.basename(image_path)} - {e}")
            self.logger.error(f"处理图片文件'{image_path}'失败.", exc_info=True)
            return False

    def add_pdf_watermark(self, input_pdf_path, filename):
        """
        使用 PyMuPDF (fitz) 为PDF文件添加页眉式水印，并原地保存。
        如果检测到已存在水印，则跳过。
        """
        if not CHINESE_FONT_AVAILABLE:
            self.log_watermark(f"[PDF] ❌ 失败: {filename} - 中文字体文件 '{CHINESE_FONT_PATH}' 未找到。")
            return False

        temp_output_path = input_pdf_path + ".tmp"
        doc = None
        try:
            doc = fitz.open(input_pdf_path)
            
            # 在处理前检查第一页是否存在水印
            if doc.page_count > 0:
                # 获取第一页的纯文本内容
                first_page_text = doc[0].get_text("text")
                # 通过一个独特的文本组合来判断水印是否存在，以降低误判率
                # 例如，我们的水印包含 "打印对象"
                if "打印对象" in first_page_text:
                    self.log_watermark(f"[PDF] ！ 跳过（已存在页眉水印）: {filename}")
                    doc.close()  # 在返回前必须关闭文档
                    doc = None 
                    return True

            total_pages = doc.page_count
            for i in range(total_pages):
                page = doc.load_page(i)
                page_rect = page.rect
                
                dynamic_fontsize = max(9, int(page_rect.width / 60))
                header_height = dynamic_fontsize * 1.8

                header_text = f"打印对象：{filename}  第 {i + 1}/{total_pages} 页"
                
                header_width = page_rect.width * 0.95
                header_left = (page_rect.width - header_width) / 2
                header_top = 5
                header_rect = fitz.Rect(
                    header_left,
                    header_top,
                    header_left + header_width,
                    header_top + header_height
                )

                page.insert_textbox(
                    header_rect,
                    header_text,
                    fontsize=dynamic_fontsize,
                    color=(1, 0, 0),
                    fontname="china-font",
                    fontfile=CHINESE_FONT_PATH,
                    align=fitz.TEXT_ALIGN_CENTER,
                )

            doc.save(temp_output_path, garbage=4, deflate=True, clean=True)
            doc.close()
            doc = None

            os.remove(input_pdf_path)
            os.rename(temp_output_path, input_pdf_path)

            self.log_watermark(f"[PDF] ✔ 成功: {filename}")
            return True

        except Exception as e:
            self.log_watermark(f"[PDF] ❌ 失败: {filename} - {e}")
            self.logger.error(f"处理PDF文件'{input_pdf_path}'失败.", exc_info=True)
            if os.path.exists(temp_output_path):
                os.remove(temp_output_path)
            return False
        finally:
            # 确保任何情况下文档都会被关闭
            if doc:
                doc.close()
    
    def start_watermark_processing(self):
        if not self.watermark_folder_path.get():
            messagebox.showwarning("警告", "请先选择一个文件夹！", parent=self.watermark_tab)
            self.logger.warning("水印任务启动失败：未选择文件夹。")
            return
        if not (self.process_word.get() or self.process_pic.get() or self.process_pdf.get() or self.process_excel.get()):
            messagebox.showwarning("警告", "请至少选择一种要处理的文件类型！", parent=self.watermark_tab)
            self.logger.warning("水印任务启动失败：未选择文件类型。")
            return
        try:
            # 只需验证透明度
            self.pic_opacity.get()
        except (tk.TclError, ValueError):
            # 更新了错误信息
            messagebox.showerror("输入错误", "图片水印的'透明度'必须是有效的数字。", parent=self.watermark_tab)
            self.logger.error("水印任务启动失败：无效的透明度。", exc_info=True)
            return
        
        if messagebox.askyesno("确认操作", "此操作将直接修改原始文件，不可撤销。\n请确保您已备份重要文件。\n\n是否继续？", parent=self.watermark_tab):
            self.logger.info("用户确认开始水印处理任务。")
            self.watermark_process_button.config(state=tk.DISABLED)
            processing_thread = threading.Thread(target=self._process_watermark_files)
            processing_thread.daemon = True
            processing_thread.start()
        else:
            self.logger.info("用户取消了水印处理任务。")
            
    def _process_watermark_files(self):
        folder = self.watermark_folder_path.get()
        self.log_watermark("\n" + "=" * 40)
        self.log_watermark("...开始处理水印任务...")
        do_word, do_excel, do_pic, do_pdf = (self.process_word.get(), self.process_excel.get(), self.process_pic.get(), self.process_pdf.get())
        self.logger.info(f"开始扫描文件夹: {folder}")
        self.logger.info(f"处理类型 - Word: {do_word}, Excel: {do_excel}, 图片: {do_pic}, PDF: {do_pdf}")
        pic_extensions = (".jpg", ".jpeg", ".png", ".bmp")
        counts = {"word": 0, "excel": 0, "pic": 0, "pdf": 0, "total": 0}
        
        for root_dir, _, files in os.walk(folder):
            for file in files:
                if file.startswith("~"):
                    continue
                filepath = os.path.join(root_dir, file)
                filename_lower = file.lower()
                if do_word and filename_lower.endswith(".docx"):
                    if self.add_word_watermark(filepath):
                        counts["word"] += 1
                    counts["total"] += 1
                elif do_excel and filename_lower.endswith(".xlsx"):
                    if self.add_excel_watermark(filepath):
                        counts["excel"] += 1
                    counts["total"] += 1
                elif do_pic and filename_lower.endswith(pic_extensions):
                    if self.add_picture_watermark(filepath):
                        counts["pic"] += 1
                    counts["total"] += 1
                elif do_pdf and filename_lower.endswith(".pdf"):
                    if self.add_pdf_watermark(filepath, file):
                        counts["pdf"] += 1
                    counts["total"] += 1
        
        summary = (
            f"\n>>> 处理完成 <<<\n"
            f"总共扫描并尝试处理 {counts['total']} 个文件。\n"
            f"  - 成功处理 Word: {counts['word']} 个\n"
            f"  - 成功处理 Excel: {counts['excel']} 个\n"
            f"  - 成功处理 图片: {counts['pic']} 个\n"
            f"  - 成功处理 PDF: {counts['pdf']} 个\n" + "=" * 40
        )
        self.log_watermark(summary)
        # 建议：使用 self.root.after 来确保线程安全
        self.root.after(0, lambda: messagebox.showinfo("处理完成", summary, parent=self.watermark_tab))
        self.root.after(0, lambda: self.watermark_process_button.config(state=tk.NORMAL))
        self.logger.info("水印处理任务完成。")

    # ==================================================================
    # II. 批量打印模块 (此部分代码未作修改)
    # ==================================================================
    # ... (后续代码与原版完全相同，此处省略以保持简洁) ...
    def _initialize_print_vars(self):
        """初始化打印标签页所需的所有变量"""
        self.print_folder_path = tk.StringVar()
        self.print_printer_name = tk.StringVar()
        self.libreoffice_path_var = tk.StringVar(value=LIBREOFFICE_PATH) 
        self.print_doc_var = tk.BooleanVar(value=True)
        self.print_docx_var = tk.BooleanVar(value=True)
        self.print_pdf_var = tk.BooleanVar(value=True)
        self.print_jpg_var = tk.BooleanVar(value=True)
        self.print_png_var = tk.BooleanVar(value=True)
        self.print_bmp_var = tk.BooleanVar(value=True)
        self.print_filter_by_pages = tk.BooleanVar(value=False)
        self.print_min_pages = tk.IntVar(value=1)
        self.print_max_pages = tk.IntVar(value=2)
        self.print_margin = tk.IntVar(value=100)

    def _setup_print_tab(self):
        """构建打印标签页的UI界面"""
        self.print_tab.grid_columnconfigure(1, weight=1)
        self.print_tab.grid_rowconfigure(8, weight=1)

        ttk.Label(self.print_tab, text="文档/图片文件夹:").grid(
            row=0, column=0, padx=5, pady=5, sticky="w"
        )
        ttk.Entry(self.print_tab, textvariable=self.print_folder_path, width=50).grid(
            row=0, column=1, padx=5, pady=5, sticky="ew"
        )
        ttk.Button(
            self.print_tab, text="选择文件夹", command=self._select_print_folder
        ).grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(self.print_tab, text="打印机名称:").grid(
            row=1, column=0, padx=5, pady=5, sticky="w"
        )
        ttk.Entry(self.print_tab, textvariable=self.print_printer_name, width=50).grid(
            row=1, column=1, padx=5, pady=5, sticky="ew"
        )
        ttk.Label(self.print_tab, text="(与系统中名称完全一致)").grid(
            row=1, column=2, padx=5, pady=5, sticky="w"
        )

        ttk.Label(self.print_tab, text="要打印的文件类型:").grid(
            row=2, column=0, padx=5, pady=5, sticky="w"
        )
        file_type_frame = ttk.Frame(self.print_tab)
        file_type_frame.grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky="w")
        ttk.Checkbutton(file_type_frame, text=".doc", variable=self.print_doc_var).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Checkbutton(
            file_type_frame, text=".docx", variable=self.print_docx_var
        ).pack(side=tk.LEFT, padx=2)
        ttk.Checkbutton(file_type_frame, text=".pdf", variable=self.print_pdf_var).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Label(file_type_frame, text="|    图片(合并为PDF打印):").pack(
            side=tk.LEFT, padx=(10, 2)
        )
        ttk.Checkbutton(file_type_frame, text=".jpg", variable=self.print_jpg_var).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Checkbutton(file_type_frame, text=".png", variable=self.print_png_var).pack(
            side=tk.LEFT, padx=2
        )
        ttk.Checkbutton(file_type_frame, text=".bmp", variable=self.print_bmp_var).pack(
            side=tk.LEFT, padx=2
        )

        ttk.Label(self.print_tab, text="页码筛选(Word/PDF):").grid(
            row=3, column=0, padx=5, pady=5, sticky="w"
        )
        filter_frame = ttk.Frame(self.print_tab)
        filter_frame.grid(row=3, column=1, columnspan=2, padx=5, pady=5, sticky="w")
        ttk.Checkbutton(
            filter_frame, text="启用筛选:", variable=self.print_filter_by_pages
        ).pack(side=tk.LEFT)
        ttk.Label(filter_frame, text="  页数 >=").pack(side=tk.LEFT)
        ttk.Spinbox(
            filter_frame, from_=0, to=9999, textvariable=self.print_min_pages, width=5
        ).pack(side=tk.LEFT)
        ttk.Label(filter_frame, text="  且  页数 <=").pack(side=tk.LEFT)
        ttk.Spinbox(
            filter_frame, from_=0, to=9999, textvariable=self.print_max_pages, width=5
        ).pack(side=tk.LEFT)

        def open_libreoffice_page(event):
            webbrowser.open_new("https://zh-cn.libreoffice.org/download/download/")
        # 创建标签并绑定事件
        label = tk.Label(self.print_tab, text="LibreOffice 安装路径（点击下载）", fg="blue", cursor="hand2")
        label.grid(row=4, column=0, padx=5, pady=5, sticky="w")
        label.bind("<Button-1>", open_libreoffice_page)

        ttk.Entry(self.print_tab, textvariable=self.libreoffice_path_var, width=50).grid(
            row=4, column=1, padx=5, pady=5, sticky="ew"
        )
        ttk.Button(
            self.print_tab, text="浏览...", command=self._select_libreoffice_path
        ).grid(row=4, column=2, padx=5, pady=5)

        ttk.Label(self.print_tab, text="图片页边距 (像素):").grid(
            row=5, column=0, padx=5, pady=5, sticky="w"
        )
        ttk.Spinbox(
            self.print_tab, from_=0, to=1000, textvariable=self.print_margin, width=10
        ).grid(row=5, column=1, padx=5, pady=5, sticky="w")

        self.print_button = ttk.Button(
            self.print_tab,
            text="开始批量打印",
            command=self.start_printing_thread,
            style="Accent.TButton",
        )
        self.print_button.grid(
            row=6, column=0, columnspan=3, padx=5, pady=10, sticky="ew", ipady=5
        )

        ttk.Label(self.print_tab, text="日志:").grid(
            row=7, column=0, padx=5, pady=5, sticky="nw"
        )
        self.print_log_area = scrolledtext.ScrolledText(
            self.print_tab,
            width=70,
            height=15,
            state=tk.DISABLED,
            font=("Consolas", 10),
        )
        self.print_log_area.grid(
            row=8, column=0, columnspan=3, padx=5, pady=5, sticky="nsew"
        )

    def log_print(self, message):
        """线程安全地向打印日志区域添加消息，并记录到文件。"""
        self.logger.info(f"[Print Tab] {message}")
        def _update():
            self.print_log_area.config(state=tk.NORMAL)
            self.print_log_area.insert(tk.END, message + "\n")
            self.print_log_area.see(tk.END)
            self.print_log_area.config(state=tk.DISABLED)
        self.root.after(0, _update)

    def _start_background_tasks(self):
        """
        启动所有不应阻塞GUI的后台初始化任务。
        目前主要是获取默认打印机。
        """
        self.log_print("正在后台检测默认打印机...")
        if PYWIN32_AVAILABLE:
            printer_thread = threading.Thread(
                target=self._fetch_default_printer_worker, 
                daemon=True
            )
            printer_thread.start()
        else:
            self.log_print(
                "【警告】'pywin32' 模块缺失，无法自动获取打印机或进行打印。"
            )
            self.log_print("请先关闭程序，运行 'pip install pywin32' 安装后再使用。")
            if hasattr(self, "print_button"):
                self.print_button.config(state=tk.DISABLED, text="依赖缺失，无法运行")
                self.logger.error("'pywin32' 模块缺失，打印功能已禁用。")

    def _fetch_default_printer_worker(self):
        """(在后台线程中运行) 负责执行获取默认打印机的耗时操作。"""
        try:
            default_printer = win32print.GetDefaultPrinter()
            self.root.after(0, self._update_printer_ui, default_printer)
        except Exception as e:
            self.root.after(0, self._handle_printer_fetch_fail, e)

    def _update_printer_ui(self, printer_name):
        """(在主GUI线程中运行) 线程安全地更新打印机名称输入框和日志。"""
        self.print_printer_name.set(printer_name)
        self.log_print(f"✔ 自动检测到默认打印机: {printer_name}")

    def _handle_printer_fetch_fail(self, error):
        """(在主GUI线程中运行) 在获取打印机失败时，线程安全地更新日志。"""
        self.log_print(f"【提示】尝试获取默认打印机失败 ({error})。请手动输入。")
        self.logger.warning("后台获取默认打印机失败。", exc_info=error)

    def _initialize_print_log(self):
        """打印模块的初始日志和检查（只包含快速操作）"""
        self.log_print("注意：勾选的图片将被合并成一个PDF文件进行打印。")
        if not os.path.exists(self.libreoffice_path_var.get()):
            self.log_print(f"【警告】未在默认路径找到LibreOffice: {self.libreoffice_path_var.get()}")
            self.log_print("【提示】如果需要打印Word(.doc/.docx)文件，请手动指定正确的'soffice.exe'路径。")


    def _select_print_folder(self):
        folder = filedialog.askdirectory(title="选择要打印的文件夹")
        if folder:
            self.print_folder_path.set(folder)
            self.log_print(f"选择的文件夹: {folder}")

    def _select_libreoffice_path(self):
        """打开文件对话框让用户选择LibreOffice (soffice.exe) 的路径。"""
        current_path = self.libreoffice_path_var.get()
        initial_dir = os.path.dirname(current_path) if os.path.exists(current_path) else r"C:\Program Files"

        filepath = filedialog.askopenfilename(
            title="选择 LibreOffice (soffice.exe)",
            initialdir=initial_dir,
            filetypes=[("Executable files", "*.exe"), ("All files", "*.*")]
        )
        if filepath:
            self.libreoffice_path_var.set(filepath)
            self.log_print(f"已更新 LibreOffice 路径: {filepath}")

    def _check_libreoffice_path(self):
        current_path = self.libreoffice_path_var.get()
        if not current_path or not os.path.exists(current_path):
            self.logger.error(f"LibreOffice路径无效或未设置: {current_path}")
            self.root.after(
                0,
                lambda: messagebox.showerror(
                    "错误",
                    f"LibreOffice 未在指定路径找到: \n{current_path}\n请在'批量打印'标签页中设置正确的路径。",
                    parent=self.print_tab,
                ),
            )
            return False
        return True

    def start_printing_thread(self):
        self.logger.info("用户点击'开始批量打印'按钮。")
        if not PYWIN32_AVAILABLE:
            messagebox.showerror(
                "依赖缺失",
                "核心功能需要的 'pywin32' 模块未安装或加载失败。\n请在命令行中运行 'pip install pywin32' 来安装它。",
                parent=self.print_tab,
            )
            self.logger.error("打印任务启动失败: pywin32 模块缺失。")
            return
        if hasattr(self, "print_button"):
            self.print_button.config(state=tk.DISABLED, text="正在处理...")
        thread = threading.Thread(target=self.run_printing_task, daemon=True)
        thread.start()

    def _windows_sort_comparator(self, a, b):
        a_base, b_base = os.path.basename(a), os.path.basename(b)
        return _StrCmpLogicalW(a_base, b_base)

    def _get_page_count(self, file_path):
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        if ext == ".pdf":
            try:
                with open(file_path, "rb") as f:
                    return len(PrintPdfReader(f).pages)
            except Exception as e:
                self.log_print(
                    f"  [页数检查] 无法读取PDF页数: {os.path.basename(file_path)}. 错误: {e}"
                )
                self.logger.warning(f"无法读取PDF '{file_path}' 的页数。", exc_info=True)
                return None
        elif ext in [".doc", ".docx"]:
            temp_dir = None
            try:
                temp_dir = tempfile.mkdtemp(prefix="page_count_")
                command = [
                    self.libreoffice_path_var.get(),
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    temp_dir,
                    file_path,
                ]
                subprocess.run(command, check=True, capture_output=True, timeout=60, creationflags=getattr(subprocess, 'CREATE_NO_WINDOW', 0))
                pdf_name = os.path.splitext(os.path.basename(file_path))[0] + ".pdf"
                temp_pdf_path = os.path.join(temp_dir, pdf_name)
                if os.path.exists(temp_pdf_path):
                    with open(temp_pdf_path, "rb") as f:
                        return len(PrintPdfReader(f).pages)
                else:
                    self.log_print(
                        f"  [页数检查] LibreOffice转换后未找到PDF文件: {pdf_name}"
                    )
                    self.logger.warning(f"LibreOffice转换后未找到PDF文件: {pdf_name}, 源文件: {file_path}")
                    return None
            except subprocess.TimeoutExpired:
                self.log_print(
                    f"  [页数检查] LibreOffice转换超时: {os.path.basename(file_path)}"
                )
                self.logger.error(f"LibreOffice转换超时: {file_path}", exc_info=True)
                return None
            except Exception as e:
                self.log_print(
                    f"  [页数检查] 获取页数失败: {os.path.basename(file_path)}. 错误: {e}"
                )
                self.logger.error(f"获取'{file_path}'页数时发生未知错误。", exc_info=True)
                return None
            finally:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
        return None

    def run_printing_task(self):
        def restore_button():
            self.root.after(
                0,
                lambda: self.print_button.config(state=tk.NORMAL, text="开始批量打印"),
            )
        
        selected_doc_types = self.print_doc_var.get() or self.print_docx_var.get()
        if selected_doc_types and not self._check_libreoffice_path():
            restore_button()
            return

        folder_path = self.print_folder_path.get()
        printer_name = self.print_printer_name.get()
        if not folder_path or not printer_name:
            msg = "请选择一个文件夹。" if not folder_path else "请输入打印机名称。"
            self.logger.warning(f"打印任务启动失败: {msg}")
            self.root.after(
                0, lambda: messagebox.showerror("错误", msg, parent=self.print_tab)
            )
            restore_button()
            return

        self.log_print("-" * 20)
        self.log_print(f"开始打印任务...")
        self.logger.info("打印任务线程已开始。")
        self.logger.info(f"文件夹: {folder_path}, 打印机: {printer_name}")
        
        if self.print_filter_by_pages.get():
            log_msg = f"页码筛选已启用: >= {self.print_min_pages.get()} 且 <= {self.print_max_pages.get()}"
            self.log_print(log_msg)
            self.logger.info(log_msg)

        image_files, other_files = [], []
        img_exts = [
            f"*.{ext}"
            for ext, var in [
                ("jpg", self.print_jpg_var),
                ("png", self.print_png_var),
                ("bmp", self.print_bmp_var),
            ]
            if var.get()
        ]
        doc_exts = [
            f"*.{ext}"
            for ext, var in [
                ("doc", self.print_doc_var),
                ("docx", self.print_docx_var),
                ("pdf", self.print_pdf_var),
            ]
            if var.get()
        ]

        for ext in img_exts:
            image_files.extend(glob.glob(os.path.join(folder_path, ext)))
        for ext in doc_exts:
            other_files.extend(glob.glob(os.path.join(folder_path, ext)))
            
        self.logger.info(f"发现 {len(image_files)} 个图片文件和 {len(other_files)} 个文档/PDF文件。")

        sorted_images = sorted(
            image_files, key=functools.cmp_to_key(self._windows_sort_comparator)
        )

        unslotted_queue = []
        temp_dir, merged_pdf_path = None, None

        try:
            if sorted_images:
                self.log_print("开始合并图片...")
                self.logger.info("开始合并图片...")
                temp_dir = tempfile.mkdtemp(prefix="batch_print_")
                a4_px_w = int((A4_WIDTH_MM / 25.4) * PRINT_DPI)
                a4_px_h = int((A4_HEIGHT_MM / 25.4) * PRINT_DPI)
                margin = self.print_margin.get()
                pdf_paths = []
                for img_path in sorted_images:
                    try:
                        with Image.open(img_path) as image:
                            if image.mode != "RGB":
                                image = image.convert("RGB")
                            draw_w, draw_h = a4_px_w - (2 * margin), a4_px_h - (
                                2 * margin
                            )
                            img_w, img_h = image.size
                            if img_w == 0 or img_h == 0:
                                self.logger.warning(f"跳过尺寸为0的图片: {img_path}")
                                continue
                            ratio = min(draw_w / img_w, draw_h / img_h)
                            new_size = (int(img_w * ratio), int(img_h * ratio))
                            resized = image.resize(new_size, Image.Resampling.LANCZOS)
                            a4_page = Image.new("RGB", (a4_px_w, a4_px_h), "white")
                            paste_x = (a4_px_w - new_size[0]) // 2
                            paste_y = (a4_px_h - new_size[1]) // 2
                            a4_page.paste(resized, (paste_x, paste_y))
                            pdf_path = os.path.join(
                                temp_dir, os.path.basename(img_path) + ".pdf"
                            )
                            a4_page.save(pdf_path, "PDF", resolution=PRINT_DPI)
                            pdf_paths.append(pdf_path)
                    except Exception as e:
                        msg = f"  转换图片失败: {os.path.basename(img_path)}. 错误: {e}"
                        self.log_print(msg)
                        self.logger.error(f"转换图片'{img_path}'失败。", exc_info=True)


                if pdf_paths:
                    merger = PrintPdfWriter()
                    for pdf in pdf_paths:
                        merger.append(pdf)
                    merged_pdf_path = os.path.join(folder_path, "_merged_images.pdf")
                    merger.write(merged_pdf_path)
                    merger.close()
                    msg = f"  图片已合并到: {os.path.basename(merged_pdf_path)}"
                    self.log_print(msg)
                    self.logger.info(msg)
                    unslotted_queue.append(merged_pdf_path)

            unslotted_queue.extend(other_files)
            if not unslotted_queue:
                self.log_print("未找到任何要打印的文件。")
                self.logger.info("未找到任何要打印的文件，任务结束。")
                self.root.after(
                    0,
                    lambda: messagebox.showinfo(
                        "完成", "未找到任何要打印的文件。", parent=self.print_tab
                    ),
                )
                restore_button()
                return

            final_print_queue = sorted(
                unslotted_queue, key=functools.cmp_to_key(self._windows_sort_comparator)
            )

            filtered_queue = []
            if self.print_filter_by_pages.get():
                self.log_print("开始执行页码筛选...")
                self.logger.info("开始执行页码筛选...")
                min_p, max_p = self.print_min_pages.get(), self.print_max_pages.get()
                for f_path in final_print_queue:
                    basename = os.path.basename(f_path)
                    if basename.lower().endswith((".doc", ".docx", ".pdf")):
                        page_count = self._get_page_count(f_path)
                        if page_count is not None and (min_p <= page_count <= max_p):
                            filtered_queue.append(f_path)
                            self.log_print(
                                f"  -> '{basename}' ({page_count}页) 符合条件，加入队列。"
                            )
                        elif page_count is None:
                            filtered_queue.append(f_path)
                            self.log_print(
                                f"  -> '{basename}' (页数未知) 默认加入队列。"
                            )
                        else:
                            self.log_print(
                                f"  -> '{basename}' ({page_count}页) 不符合条件，已跳过。"
                            )
                    else:
                        filtered_queue.append(f_path)
            else:
                filtered_queue = final_print_queue

            self.log_print("-" * 20)
            self.log_print("最终待打印文件列表:")
            self.logger.info(f"最终待打印文件列表 ({len(filtered_queue)} 个):")
            
            if not filtered_queue:
                self.log_print("  (无文件通过筛选)")
                self.logger.info("  (无文件通过筛选)")
            for i, f in enumerate(filtered_queue):
                log_line = f"  {i+1}. {os.path.basename(f)}"
                self.log_print(log_line)
                self.logger.info(log_line)
            self.log_print("-" * 20)
            self.logger.info("-" * 20)


            if not filtered_queue:
                self.log_print("没有文件需要打印。")
                self.logger.info("没有文件需要打印，任务结束。")
                self.root.after(
                    0,
                    lambda: messagebox.showinfo(
                        "完成", "没有文件满足筛选条件。", parent=self.print_tab
                    ),
                )
                restore_button()
                return

            success, fail = 0, 0
            for file_path in filtered_queue:
                self.log_print(f"正在打印: {os.path.basename(file_path)}")
                self.logger.info(f"提交打印任务 for: {file_path}")
                try:
                    command = [
                        self.libreoffice_path_var.get(),
                        "--headless",
                        "--pt",
                        printer_name,
                        file_path,
                    ]
                    creationflags = 0
                    if os.name == 'nt':
                        creationflags = subprocess.CREATE_NO_WINDOW
                        
                    result = subprocess.run(
                        command,
                        capture_output=True,
                        text=True,
                        check=True,
                        creationflags=creationflags,
                        encoding='utf-8',
                        errors='ignore'
                    )
                    self.log_print(f"  ✔ 成功发送到打印机。")
                    self.logger.info(f"成功打印: {file_path}")
                    success += 1
                except subprocess.CalledProcessError as e:
                    error_msg = f"  ❌ 打印失败: {e.stderr.strip() or e.stdout.strip()}"
                    self.log_print(error_msg)
                    self.logger.error(f"打印失败. 文件: {file_path}. 返回码: {e.returncode}. STDOUT: {e.stdout}. STDERR: {e.stderr}", exc_info=False)
                    fail += 1
                except Exception as e:
                    error_msg = f"  ❌ 发生未知错误: {e}"
                    self.log_print(error_msg)
                    self.logger.error(f"打印时发生未知错误. 文件: {file_path}", exc_info=True)
                    fail += 1

            summary = f"打印完成！\n成功: {success}\n失败: {fail}"
            self.log_print(summary)
            self.logger.info(f"打印任务完成。成功: {success}, 失败: {fail}")
            self.root.after(
                0, lambda: messagebox.showinfo("完成", summary, parent=self.print_tab)
            )

        finally:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                self.log_print("已清理临时目录。")
                self.logger.info("已清理打印任务的临时目录。")
            restore_button()


if __name__ == "__main__":
    root = tk.Tk()
    # 初始隐藏窗口，避免闪烁
    root.withdraw() 
    app = PrintALLApp(root)
    # 初始化完成后再显示窗口
    root.deiconify() 
    root.mainloop()
